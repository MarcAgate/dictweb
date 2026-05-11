#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sqlite3
import traceback
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml

from tibetan_sort.tibetan_sort import TibetanSort
from pyewts import pyewts


BASE_DIR = Path(__file__).resolve().parent
DB_PATH = BASE_DIR / "MsTibTool.db"
OUTPUT_FILE = "Glossaire_Tibetain_3_Lignes_BDRC1_def.docx"

TIB_START = 0x0F00
TIB_END = 0x0FFF


def connect_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def split_tibetan_segments(text):
    segments = []
    if not text:
        return segments

    current = []
    current_is_tib = None

    for ch in text:
        code = ord(ch)
        is_tib = TIB_START <= code <= TIB_END

        if current_is_tib is None:
            current_is_tib = is_tib
            current.append(ch)
        elif is_tib == current_is_tib:
            current.append(ch)
        else:
            segments.append(("".join(current), current_is_tib))
            current = [ch]
            current_is_tib = is_tib

    if current:
        segments.append(("".join(current), current_is_tib))

    return segments


def set_tibetan_font(run):
    font = run.font
    font.name = "Microsoft Himalaya"
    font.size = Pt(16)

    r = run._element
    rPr = r.get_or_add_rPr()
    if rPr.rFonts is not None:
        rFonts = rPr.rFonts
    else:
        rFonts = rPr.get_or_add_rFonts()

    rFonts.set(qn("w:eastAsia"), "Microsoft Himalaya")


def add_text_with_tibetan_font(paragraph, text):
    for segment, is_tib in split_tibetan_segments(text or ""):
        run = paragraph.add_run(segment)
        if is_tib:
            set_tibetan_font(run)


def shade_cell(cell, fill="D9EAF7"):
    shading_elm = parse_xml(
        rf'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{fill}" />'
    )
    cell._element.get_or_add_tcPr().append(shading_elm)


def normalize_extracted_text(text):
    if not text:
        return ""

    lines = []
    for line in text.split("\n"):
        cleaned = " ".join(line.replace("\xa0", " ").split()).strip(" ,;")
        if cleaned:
            lines.append(cleaned)

    return "\n".join(lines).strip()


def extract_section_text(full_text, section_name, next_sections):
    lower_full = full_text.lower()
    section_label = f"{section_name.lower()} :"
    start = lower_full.find(section_label)

    if start == -1:
        return ""

    content_start = start + len(section_label)
    end = len(full_text)

    for next_name in next_sections:
        next_label = f"{next_name.lower()} :"
        idx = lower_full.find(next_label, content_start)
        if idx != -1 and idx < end:
            end = idx

    extracted = full_text[content_start:end].strip()
    return normalize_extracted_text(extracted)


def extract_sens_et_variantes(html_def):
    if not html_def:
        return []

    soup = BeautifulSoup(html_def, "html.parser")

    for br in soup.find_all("br"):
        br.replace_with("\n")

    full_text = soup.get_text("\n")
    full_text = full_text.replace("\r\n", "\n").replace("\r", "\n")

    sens = extract_section_text(
        full_text,
        "Sens racine",
        ["Variantes", "Définition", "Definition", "Étymologie", "Etymologie"],
    )

    variantes = extract_section_text(
        full_text,
        "Variantes",
        ["Définition", "Definition", "Étymologie", "Etymologie"],
    )

    blocks = []
    if sens:
        blocks.append(("Sens racine :", sens))
    if variantes:
        blocks.append(("Variantes :", variantes))

    return blocks


def create_word_document(rows):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    title = doc.add_heading("Glossaire Tibétain Bouddhique", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Glossaire Encyclopédique (Classement BDRC)")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    for row in rows:
        rec_id = row["id"]
        tib = row["tib"] or ""
        wylie = row["wylie"] or ""
        def_html = row["def"] or ""

        entry_title = doc.add_heading(level=2)
        entry_title_run = entry_title.add_run(f"[ID: {rec_id}] ")
        entry_title_run.font.size = Pt(11)
        entry_title_run.font.bold = True
        add_text_with_tibetan_font(entry_title, tib)
        entry_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

        table = doc.add_table(rows=2, cols=2)
        table.style = "Light Grid Accent 1"
        table.autofit = False
        table.allow_autofit = False

        col1_width = Inches(2.5)
        col2_width = Inches(7.5)

        table.columns[0].width = col1_width
        table.columns[1].width = col2_width

        for table_row in table.rows:
            table_row.cells[0].width = col1_width
            table_row.cells[1].width = col2_width

        # Ligne 1 : Tibétain / Wylie
        row1 = table.rows[0].cells
        shade_cell(row1[0])

        row1[0].text = ""
        p_label = row1[0].paragraphs[0]
        run_label = p_label.add_run("Tibétain / Wylie")
        run_label.font.bold = True
        run_label.font.size = Pt(10)
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER

        row1[1].text = ""
        p_tib = row1[1].paragraphs[0]
        add_text_with_tibetan_font(p_tib, tib)

        if wylie.strip():
            p_wylie = row1[1].add_paragraph()
            run_wylie = p_wylie.add_run(wylie)
            run_wylie.font.size = Pt(10)

        # Ligne 2 : Significations
        row2 = table.rows[1].cells
        shade_cell(row2[0])

        row2[0].text = ""
        p_label = row2[0].paragraphs[0]
        run_label = p_label.add_run("Significations")
        run_label.font.bold = True
        run_label.font.size = Pt(10)
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER

        row2[1].text = ""
        signification_blocks = extract_sens_et_variantes(def_html)

        if signification_blocks:
            first_block = True
            for block_label, block_text in signification_blocks:
                if first_block:
                    p_block_label = row2[1].paragraphs[0]
                    first_block = False
                else:
                    row2[1].add_paragraph("")
                    p_block_label = row2[1].add_paragraph()

                run_block_label = p_block_label.add_run(block_label)
                run_block_label.font.bold = True
                run_block_label.font.size = Pt(10)

                for line in block_text.split("\n"):
                    p_line = row2[1].add_paragraph()
                    add_text_with_tibetan_font(p_line, line)
        else:
            row2[1].paragraphs[0].add_run("—")

        doc.add_paragraph()

    return doc


def sort_rows_by_tibetan(rows):
    converter = pyewts()
    sorter = TibetanSort()

    seen_ids = set()
    unique_rows = []
    for row in rows:
        row_id = row["id"]
        if row_id not in seen_ids:
            unique_rows.append(row)
            seen_ids.add(row_id)

    print(f"Déduplication : {len(rows)} → {len(unique_rows)} entrées uniques")

    def get_tibetan_key(row):
        tib = (row["tib"] or "").strip()
        wylie = (row["wylie"] or "").strip()

        if tib:
            return tib
        if wylie:
            try:
                return converter.toUnicode(wylie)
            except Exception:
                return wylie
        return ""

    rows_with_keys = []
    for row in unique_rows:
        tibetan_key = get_tibetan_key(row)
        rows_with_keys.append((tibetan_key, row))

    all_tibetan_keys = [item[0] for item in rows_with_keys]
    sorted_tibetan_keys = sorter.sort_list(all_tibetan_keys)
    key_to_position = {key: i for i, key in enumerate(sorted_tibetan_keys)}

    rows_with_keys_sorted = sorted(
        rows_with_keys,
        key=lambda x: key_to_position.get(x[0], 999999)
    )

    sorted_rows = [row for tibetan_key, row in rows_with_keys_sorted]

    print(f"Tri appliqué : {len(sorted_rows)} entrées")
    return sorted_rows


def fetch_rows():
    conn = connect_db()
    try:
        cur = conn.cursor()
        cur.execute("""
            SELECT id, tib, wylie, def
            FROM dict
            WHERE source = 'RMY'
            AND contexte = 'Glossaire de l''Encyclopédie'
            ORDER BY id
            
        """)
        return cur.fetchall()
    finally:
        conn.close()


def main():
    try:
        print("📊 Récupération des données de dict...")
        rows = fetch_rows()
        print(f"{len(rows)} entrées trouvées (avant déduplication)")

        print("🔤 Tri selon l'ordre alphabétique tibétain BDRC...")
        rows_sorted = sort_rows_by_tibetan(rows)
        print(f"✓ Tri appliqué ({len(rows_sorted)} entrées uniques)")

        print("📝 Génération du document Word à partir du champ def...")
        doc = create_word_document(rows_sorted)
        doc.save(OUTPUT_FILE)

        print(f"\n✅ Document généré avec succès : {OUTPUT_FILE}")
        print(f"{len(rows_sorted)} entrées triées en ordre tibétain conforme BDRC")

    except Exception:
        print("❌ Erreur lors de la génération du document :")
        traceback.print_exc()


if __name__ == "__main__":
    main()