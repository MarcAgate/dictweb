#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import mariadb
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import traceback

from tibetan_sort.tibetan_sort import TibetanSort
from pyewts import pyewts

DB_CONFIG = {
    "host": "localhost",
    "port": 3306,
    "user": "root",
    "password": "illu2025!",
    "database": "illu-eng",
}

OUTPUT_FILE = "Glossaire_Tibetain_3_Lignes_BDRC1.docx"

TIB_START = 0x0F00
TIB_END = 0x0FFF


def connect_db():
    return mariadb.connect(
        host=DB_CONFIG["host"],
        port=DB_CONFIG["port"],
        user=DB_CONFIG["user"],
        password=DB_CONFIG["password"],
        database=DB_CONFIG["database"],
    )


def split_tibetan_segments(text):
    segments = []
    if not text:
        return segments

    current = []
    current_is_tib = None

    for ch in text:
        code = ord(ch)
        is_tib = TIB_START <= code <= TIB_END

        if current_is_tib is None:
            current_is_tib = is_tib
            current.append(ch)
        elif is_tib == current_is_tib:
            current.append(ch)
        else:
            segments.append(("".join(current), current_is_tib))
            current = [ch]
            current_is_tib = is_tib

    if current:
        segments.append(("".join(current), current_is_tib))

    return segments


def set_tibetan_font(run):
    font = run.font
    font.name = "Microsoft Himalaya"
    font.size = Pt(16)

    r = run._element
    rPr = r.get_or_add_rPr()
    if rPr.rFonts is not None:
        rFonts = rPr.rFonts
    else:
        rFonts = rPr.get_or_add_rFonts()

    rFonts.set(qn("w:eastAsia"), "Microsoft Himalaya")


def add_text_with_tibetan_font(paragraph, text):
    for segment, is_tib in split_tibetan_segments(text or ""):
        run = paragraph.add_run(segment)
        if is_tib:
            set_tibetan_font(run)


def shade_cell(cell, fill="D9EAF7"):
    shading_elm = parse_xml(
        rf'<w:shd w:fill="{fill}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )
    cell._element.get_or_add_tcPr().append(shading_elm)


def create_word_document(rows):
    """
    rows = [
        (id, tib, wylie, etymologie, sensracine,
         tdctib, tdcfr, iadef, rmydef, illfr),
        ...
    ]

    Structure fixe :
    Ligne 1 : Tibétain / Wylie
    Ligne 2 : Significations
    Ligne 3 : Définition tibétaine / Traduction
    """
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    title = doc.add_heading("Glossaire Tibétain Bouddhique", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Glossaire Encyclopédique (Classement BDRC)")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    for row in rows:
        rec_id, tib, wylie, etymologie, sensracine, tdctib, tdcfr, iadef, rmydef, illfr = row

        entry_title = doc.add_heading(level=2)
        entry_title_run = entry_title.add_run(f"[ID: {rec_id}] ")
        entry_title_run.font.size = Pt(11)
        entry_title_run.font.bold = True
        add_text_with_tibetan_font(entry_title, tib or "")
        entry_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

        table = doc.add_table(rows=3, cols=2)
        table.style = "Light Grid Accent 1"
        table.autofit = False
        table.allow_autofit = False

        col1_width = Inches(2.5)   # 1/4
        col2_width = Inches(7.5)   # 3/4

        table.columns[0].width = col1_width
        table.columns[1].width = col2_width

        for row in table.rows:
            row.cells[0].width = col1_width
            row.cells[1].width = col2_width

        # ─────────────────────────────────────────────
        # Ligne 1 : Tibétain / Wylie
        # ─────────────────────────────────────────────
        row1 = table.rows[0].cells
        shade_cell(row1[0])

        row1[0].text = ""
        p_label = row1[0].paragraphs[0]
        run_label = p_label.add_run("Tibétain / Wylie")
        run_label.font.bold = True
        run_label.font.size = Pt(10)
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER

        row1[1].text = ""
        p_tib = row1[1].paragraphs[0]
        add_text_with_tibetan_font(p_tib, tib or "")

        if wylie and wylie.strip():
            p_wylie = row1[1].add_paragraph()
            run_wylie = p_wylie.add_run(wylie)
            run_wylie.font.size = Pt(10)

        # ─────────────────────────────────────────────
        # Ligne 2 : Significations
        # ─────────────────────────────────────────────
        row2 = table.rows[1].cells
        shade_cell(row2[0])

        row2[0].text = ""
        p_label = row2[0].paragraphs[0]
        run_label = p_label.add_run("Significations")
        run_label.font.bold = True
        run_label.font.size = Pt(10)
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER

        row2[1].text = ""
        has_sens = bool(sensracine and sensracine.strip())
        has_nuances = bool(rmydef and rmydef.strip())

        if has_sens:
            p_sr_label = row2[1].paragraphs[0]
            run_sr_label = p_sr_label.add_run("Sens racine :")
            run_sr_label.font.bold = True
            run_sr_label.font.size = Pt(10)

            p_sr = row2[1].add_paragraph()
            add_text_with_tibetan_font(p_sr, sensracine)

        if has_nuances:
            if has_sens:
                row2[1].add_paragraph("")
            p_nu = row2[1].add_paragraph()
            run_nu = p_nu.add_run("Nuances :")
            run_nu.font.bold = True
            run_nu.font.size = Pt(10)

            p_nu_text = row2[1].add_paragraph()
            add_text_with_tibetan_font(p_nu_text, rmydef)

        if not has_sens and not has_nuances:
            p_empty = row2[1].paragraphs[0]
            p_empty.add_run("—")

        # ─────────────────────────────────────────────
        # Ligne 3 : Définition tibétaine / Traduction
        # ─────────────────────────────────────────────
        row3 = table.rows[2].cells
        shade_cell(row3[0])

        row3[0].text = ""
        p_label = row3[0].paragraphs[0]
        run_label = p_label.add_run("Définition tibétaine /\nTraduction")
        run_label.font.bold = True
        run_label.font.size = Pt(9)
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER

        row3[1].text = ""
        has_tdctib = bool(tdctib and tdctib.strip())
        has_tdcfr = bool(tdcfr and tdcfr.strip())

        if has_tdctib:
            p_tdctib = row3[1].paragraphs[0]
            add_text_with_tibetan_font(p_tdctib, tdctib)

        if has_tdcfr:
            if has_tdctib:
                row3[1].add_paragraph("")
            p_trad_label = row3[1].add_paragraph()
            run_trad = p_trad_label.add_run("Traduction :")
            run_trad.font.bold = True
            run_trad.font.size = Pt(10)

            p_tdcfr = row3[1].add_paragraph()
            add_text_with_tibetan_font(p_tdcfr, tdcfr)

        if not has_tdctib and not has_tdcfr:
            p_empty = row3[1].paragraphs[0]
            p_empty.add_run("—")

        doc.add_paragraph()

    return doc


def sort_rows_by_tibetan(rows):
    """
    Tri selon l'ordre tibétain BDRC sans doublons
    """
    converter = pyewts()
    sorter = TibetanSort()

    seen_ids = set()
    unique_rows = []
    for row in rows:
        row_id = row[0]
        if row_id not in seen_ids:
            unique_rows.append(row)
            seen_ids.add(row_id)

    print(f"Déduplication : {len(rows)} → {len(unique_rows)} entrées uniques")

    def get_tibetan_key(row):
        tib = row[1]
        wylie = row[2]

        if tib and tib.strip():
            return tib
        if wylie and wylie.strip():
            try:
                return converter.toUnicode(wylie)
            except Exception:
                return wylie
        return ""

    rows_with_keys = []
    for row in unique_rows:
        tibetan_key = get_tibetan_key(row)
        rows_with_keys.append((tibetan_key, row))

    all_tibetan_keys = [item[0] for item in rows_with_keys]
    sorted_tibetan_keys = sorter.sort_list(all_tibetan_keys)

    key_to_position = {key: i for i, key in enumerate(sorted_tibetan_keys)}

    rows_with_keys_sorted = sorted(
        rows_with_keys,
        key=lambda x: key_to_position.get(x[0], 999999)
    )

    sorted_rows = [row for tibetan_key, row in rows_with_keys_sorted]

    print(f"Tri appliqué : {len(sorted_rows)} entrées")
    return sorted_rows


def fetch_rows():
    conn = connect_db()
    try:
        cur = conn.cursor()
        query = """
        SELECT
            id,
            tib,
            wylie,
            etymologie,
            sensracine,
            tdctib,
            tdcfr,
            iadef,
            rmydef,
            illfr
        FROM glossenc
        ORDER BY id
        """
        cur.execute(query)
        rows = cur.fetchall()
        return rows
    finally:
        conn.close()


def main():
    try:
        print("📊 Récupération des données de glossenc...")
        rows = fetch_rows()
        print(f"{len(rows)} entrées trouvées (avant déduplication)")

        print("🔤 Tri selon l'ordre alphabétique tibétain BDRC...")
        rows_sorted = sort_rows_by_tibetan(rows)
        print(f"✓ Tri appliqué ({len(rows_sorted)} entrées uniques)")

        print("📝 Génération du document Word...")
        doc = create_word_document(rows_sorted)
        doc.save(OUTPUT_FILE)

        print(f"\n✅ Document généré avec succès : {OUTPUT_FILE}")
        print(f"{len(rows_sorted)} entrées triées en ordre tibétain conforme BDRC")

    except Exception:
        print("❌ Erreur lors de la génération du document :")
        traceback.print_exc()


if __name__ == "__main__":
    main()